import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
import time
import textwrap
import json

from database import get_user, init_db, save_history, get_history, delete_user_account
from auth import register, login, reset_password
from otp import (
    generate_otp, verify_otp, get_otp_for_display,
    is_fp_locked, record_fp_failure, reset_fp_attempts, fp_attempts_left
)
from utils import (
    analyze_dataframe, generate_summary,
    generate_quick_scan, generate_executive_summary,
    generate_full_report, answer_user_question, numeric_stats
)

# ──────────────────────────────────────────────
#  INIT
# ──────────────────────────────────────────────
init_db()

st.set_page_config(
    page_title="BizIntel AI",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ──────────────────────────────────────────────
#  SESSION STATE
# ──────────────────────────────────────────────
defaults = {
    "logged_in":       False,
    "page":            "login",
    "user":            None,
    "first_visit":     True,
    "df_analyzed":     None,
    "report_text":     "",
    "summary":         None,
    "filename":        "",
    "selected_column": None,
    "business_goal":   "",
    "chat_history":    [],
    "all_figs":        {},
    "report_type":     "Full Report",
    "graph_types":     ["Bar Chart"],
    "theme":           "Dark",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ──────────────────────────────────────────────
#  THEME CSS
# ──────────────────────────────────────────────
def apply_theme():
    dark = st.session_state.theme == "Dark"
    bg        = "#0d1117"   if dark else "#f0f4f8"
    bg2       = "#1a2035"   if dark else "#ffffff"
    bg3       = "#1e2433"   if dark else "#e8edf3"
    border    = "#2d3748"   if dark else "#cbd5e0"
    text      = "#e2e8f0"   if dark else "#1a202c"
    subtext   = "#8892a4"   if dark else "#4a5568"
    accent    = "#00d4ff"   if dark else "#0066cc"
    banner_bg = "linear-gradient(135deg,#0f3460,#16213e)" if dark else "linear-gradient(135deg,#dbeafe,#eff6ff)"
    btn_bg    = "linear-gradient(135deg,#0f3460,#00d4ff20)" if dark else "linear-gradient(135deg,#dbeafe,#bfdbfe)"
    report_bg = "#1a2035"   if dark else "#ffffff"

    st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
    html, body, [class*="css"] {{ font-family: 'Inter', sans-serif; }}
    .main {{ background: {bg}; }}

    .metric-card {{
        background: {bg3};
        border: 1px solid {border};
        border-radius: 12px;
        padding: 20px;
        text-align: center;
        margin-bottom: 10px;
    }}
    .metric-value {{ font-size: 2rem; font-weight: 700; }}
    .metric-label {{ font-size: 0.85rem; color: {subtext}; margin-top: 4px; }}

    .welcome-banner {{
        background: {banner_bg};
        border-left: 4px solid {accent};
        border-radius: 8px;
        padding: 16px 24px;
        margin-bottom: 24px;
    }}
    .welcome-banner h2 {{ color: {accent}; margin: 0; }}
    .welcome-banner p  {{ color: {subtext}; margin: 4px 0 0; }}

    .section-header {{
        border-left: 3px solid {accent};
        padding-left: 12px;
        margin: 24px 0 12px;
        font-weight: 600;
        color: {text};
    }}

    div[data-testid="stDownloadButton"] button {{
        width: 100%;
        background: {btn_bg};
        border: 1px solid {accent};
        color: {accent};
        border-radius: 8px;
        font-weight: 600;
    }}

    .history-card {{
        background: {bg2};
        border: 1px solid {border};
        border-radius: 10px;
        padding: 14px 18px;
        margin-bottom: 10px;
    }}
    .history-meta {{ font-size: 0.78rem; color: {subtext}; }}

    .report-box {{
        background: {report_bg};
        border: 1px solid {border};
        border-radius: 10px;
        padding: 20px 24px;
        white-space: pre-wrap;
        font-size: 0.88rem;
        line-height: 1.75;
        color: {text};
        max-height: 620px;
        overflow-y: auto;
        font-family: 'Courier New', monospace;
    }}

    .chat-bubble-user {{
        background: {"#1e3a5f" if dark else "#dbeafe"};
        border: 1px solid {"#2d5a8e" if dark else "#93c5fd"};
        border-radius: 12px 12px 2px 12px;
        padding: 12px 16px;
        margin: 8px 0 8px 40px;
        color: {text};
        font-size: 0.9rem;
    }}
    .chat-bubble-ai {{
        background: {bg2};
        border: 1px solid {border};
        border-radius: 2px 12px 12px 12px;
        padding: 12px 16px;
        margin: 8px 40px 8px 0;
        color: {text};
        font-size: 0.9rem;
        white-space: pre-wrap;
    }}

    .stTabs [data-baseweb="tab"] {{ font-weight: 600; }}
    .theme-toggle-bar {{
        background: {bg3};
        border: 1px solid {border};
        border-radius: 8px;
        padding: 6px 16px;
        margin-bottom: 12px;
        display: flex;
        align-items: center;
        gap: 12px;
    }}
</style>
""", unsafe_allow_html=True)

apply_theme()

# ──────────────────────────────────────────────
#  CHART CONFIG (adapts to theme)
# ──────────────────────────────────────────────
def chart_colors():
    dark = st.session_state.theme == "Dark"
    return {
        "PALETTE":   {"Positive": "#22c55e", "Negative": "#ef4444", "Neutral": "#f59e0b"},
        "FACECOLOR": "#0d1117" if dark else "#ffffff",
        "TEXTCOLOR": "#e2e8f0" if dark else "#1a202c",
        "GRIDCOLOR": "#1f2937" if dark else "#e2e8f0",
        "ACCENT":    "#00d4ff" if dark else "#0066cc",
    }


# ──────────────────────────────────────────────
#  CHART BUILDERS
# ──────────────────────────────────────────────
def build_bar_chart(df):
    cc = chart_colors()
    counts = df["Sentiment"].value_counts()
    fig, ax = plt.subplots(figsize=(7, 4), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    bars = ax.bar(counts.index, counts.values,
                  color=[cc["PALETTE"].get(s, "#888") for s in counts.index],
                  edgecolor=cc["GRIDCOLOR"], linewidth=0.5, width=0.5)
    for bar in bars:
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                str(int(bar.get_height())), ha="center", va="bottom",
                color=cc["TEXTCOLOR"], fontsize=10, fontweight="bold")
    ax.set_title("Sentiment Distribution", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.set_xlabel("Sentiment", color=cc["TEXTCOLOR"]); ax.set_ylabel("Count", color=cc["TEXTCOLOR"])
    ax.tick_params(colors=cc["TEXTCOLOR"]); ax.spines[:].set_color(cc["GRIDCOLOR"])
    ax.yaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.5)
    fig.tight_layout(); return fig


def build_pie_chart(df):
    cc = chart_colors()
    counts = df["Sentiment"].value_counts()
    fig, ax = plt.subplots(figsize=(6, 5), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    wedge_colors = [cc["PALETTE"].get(s, "#888") for s in counts.index]
    wedges, texts, autotexts = ax.pie(
        counts.values, labels=counts.index, autopct="%1.1f%%",
        colors=wedge_colors, startangle=90,
        wedgeprops={"edgecolor": cc["FACECOLOR"], "linewidth": 2},
        textprops={"color": cc["TEXTCOLOR"]}
    )
    for at in autotexts:
        at.set_fontsize(10); at.set_fontweight("bold")
    ax.set_title("Sentiment Share", color=cc["TEXTCOLOR"], fontsize=13)
    fig.tight_layout(); return fig


def build_histogram(df):
    cc = chart_colors()
    fig, ax = plt.subplots(figsize=(7, 4), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    ax.hist(df["Polarity"], bins=30, color=cc["ACCENT"],
            edgecolor=cc["FACECOLOR"], linewidth=0.4, alpha=0.85)
    ax.axvline(0, color="#ef4444", linewidth=1.5, linestyle="--", label="Neutral boundary")
    ax.axvline(df["Polarity"].mean(), color="#22c55e", linewidth=1.5,
               linestyle="-.", label=f"Mean ({df['Polarity'].mean():.2f})")
    ax.set_title("Polarity Score Distribution", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.set_xlabel("Polarity Score", color=cc["TEXTCOLOR"]); ax.set_ylabel("Frequency", color=cc["TEXTCOLOR"])
    ax.tick_params(colors=cc["TEXTCOLOR"]); ax.spines[:].set_color(cc["GRIDCOLOR"])
    ax.yaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.5)
    ax.legend(facecolor=cc["FACECOLOR"], labelcolor=cc["TEXTCOLOR"], fontsize=8)
    fig.tight_layout(); return fig


def build_scatter(df):
    cc = chart_colors()
    fig, ax = plt.subplots(figsize=(7, 4), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    for sentiment, color in cc["PALETTE"].items():
        sub = df[df["Sentiment"] == sentiment]
        ax.scatter(sub["Polarity"], sub["Subjectivity"],
                   c=color, label=sentiment, alpha=0.6, s=20, edgecolors="none")
    ax.set_title("Polarity vs Subjectivity", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.set_xlabel("Polarity", color=cc["TEXTCOLOR"]); ax.set_ylabel("Subjectivity", color=cc["TEXTCOLOR"])
    ax.tick_params(colors=cc["TEXTCOLOR"]); ax.spines[:].set_color(cc["GRIDCOLOR"])
    ax.xaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.4)
    ax.yaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.4)
    ax.legend(facecolor=cc["FACECOLOR"], labelcolor=cc["TEXTCOLOR"], fontsize=8)
    fig.tight_layout(); return fig


def build_line_chart(df):
    cc = chart_colors()
    fig, ax = plt.subplots(figsize=(7, 4), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    df_reset = df.reset_index(drop=True)
    window = max(1, len(df_reset) // 20)
    rolling_pol = df_reset["Polarity"].rolling(window=window, min_periods=1).mean()
    ax.plot(df_reset.index, rolling_pol, color=cc["ACCENT"],
            linewidth=1.8, label=f"Polarity (rolling w={window})")
    ax.axhline(0, color="#ef4444", linewidth=1, linestyle="--")
    ax.fill_between(df_reset.index, rolling_pol, 0,
                    where=(rolling_pol > 0), alpha=0.15, color="#22c55e")
    ax.fill_between(df_reset.index, rolling_pol, 0,
                    where=(rolling_pol < 0), alpha=0.15, color="#ef4444")
    ax.set_title("Polarity Trend Over Records", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.set_xlabel("Record Index", color=cc["TEXTCOLOR"]); ax.set_ylabel("Polarity", color=cc["TEXTCOLOR"])
    ax.tick_params(colors=cc["TEXTCOLOR"]); ax.spines[:].set_color(cc["GRIDCOLOR"])
    ax.yaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.4)
    ax.legend(facecolor=cc["FACECOLOR"], labelcolor=cc["TEXTCOLOR"], fontsize=8)
    fig.tight_layout(); return fig


def build_boxplot(df):
    cc = chart_colors()
    fig, ax = plt.subplots(figsize=(7, 4), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    sentiments = ["Positive", "Negative", "Neutral"]
    data_box   = [df[df["Sentiment"] == s]["Polarity"].values for s in sentiments if s in df["Sentiment"].values]
    labels_box = [s for s in sentiments if s in df["Sentiment"].values]
    bp = ax.boxplot(data_box, labels=labels_box, patch_artist=True,
                    medianprops={"color": cc["ACCENT"], "linewidth": 2})
    for patch, sent in zip(bp["boxes"], labels_box):
        patch.set_facecolor(cc["PALETTE"].get(sent, "#888") + "55")
        patch.set_edgecolor(cc["PALETTE"].get(sent, "#888"))
    for whisker in bp["whiskers"]: whisker.set_color(cc["TEXTCOLOR"])
    for cap in bp["caps"]: cap.set_color(cc["TEXTCOLOR"])
    ax.set_title("Polarity Distribution by Sentiment", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.set_ylabel("Polarity Score", color=cc["TEXTCOLOR"])
    ax.tick_params(colors=cc["TEXTCOLOR"]); ax.spines[:].set_color(cc["GRIDCOLOR"])
    ax.yaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.4)
    fig.tight_layout(); return fig


def build_heatmap(df):
    cc = chart_colors()
    numeric_df = df.select_dtypes(include=[np.number])
    numeric_df = numeric_df.drop(columns=[c for c in ["Polarity", "Subjectivity"] if c in numeric_df.columns], errors="ignore")
    if len(numeric_df.columns) < 2:
        return None
    fig, ax = plt.subplots(figsize=(8, 5), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    corr = numeric_df.corr()
    sns.heatmap(corr, ax=ax, cmap="coolwarm", annot=True, fmt=".2f",
                linewidths=0.5, linecolor=cc["FACECOLOR"],
                annot_kws={"color": cc["TEXTCOLOR"], "size": 9},
                cbar_kws={"shrink": 0.8})
    ax.set_title("Correlation Heatmap", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.tick_params(colors=cc["TEXTCOLOR"])
    fig.tight_layout(); return fig


def build_wordcount(df, text_col):
    cc = chart_colors()
    df = df.copy()
    df["word_count"] = df[text_col].fillna("").apply(lambda x: len(str(x).split()))
    fig, ax = plt.subplots(figsize=(7, 4), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    for s, color in cc["PALETTE"].items():
        subset = df[df["Sentiment"] == s]["word_count"]
        if len(subset) > 0:
            ax.hist(subset, bins=20, color=color, alpha=0.6, label=s, edgecolor=cc["FACECOLOR"])
    ax.set_title("Word Count Distribution by Sentiment", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.set_xlabel("Word Count", color=cc["TEXTCOLOR"]); ax.set_ylabel("Frequency", color=cc["TEXTCOLOR"])
    ax.tick_params(colors=cc["TEXTCOLOR"]); ax.spines[:].set_color(cc["GRIDCOLOR"])
    ax.yaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.4)
    ax.legend(facecolor=cc["FACECOLOR"], labelcolor=cc["TEXTCOLOR"], fontsize=8)
    fig.tight_layout(); return fig


def build_numeric_chart(df):
    """Bar chart of means for all numeric columns (excluding Polarity/Subjectivity)."""
    cc = chart_colors()
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    num_cols = [c for c in num_cols if c not in {"Polarity", "Subjectivity"}]
    if not num_cols:
        return None
    means = [df[c].mean() for c in num_cols]
    fig, ax = plt.subplots(figsize=(max(7, len(num_cols) * 1.2), 4), facecolor=cc["FACECOLOR"])
    ax.set_facecolor(cc["FACECOLOR"])
    bars = ax.bar(num_cols, means, color=cc["ACCENT"], edgecolor=cc["GRIDCOLOR"], linewidth=0.5, width=0.5)
    for bar, val in zip(bars, means):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(),
                f"{val:.1f}", ha="center", va="bottom",
                color=cc["TEXTCOLOR"], fontsize=8, fontweight="bold")
    ax.set_title("Numeric Column Means", color=cc["TEXTCOLOR"], fontsize=13, pad=12)
    ax.set_ylabel("Mean Value", color=cc["TEXTCOLOR"])
    ax.tick_params(colors=cc["TEXTCOLOR"], axis="x", rotation=30)
    ax.spines[:].set_color(cc["GRIDCOLOR"])
    ax.yaxis.grid(True, color=cc["GRIDCOLOR"], linewidth=0.5)
    fig.tight_layout(); return fig


CHART_BUILDERS = {
    "Bar Chart":               build_bar_chart,
    "Pie Chart":               build_pie_chart,
    "Histogram":               build_histogram,
    "Scatter Plot":            build_scatter,
    "Line Chart":              build_line_chart,
    "Box Plot":                build_boxplot,
    "Heatmap":                 build_heatmap,
    "Numeric Column Means":    build_numeric_chart,
    "Word Count Distribution": None,   # handled separately
}


# ──────────────────────────────────────────────
#  FILE LOADER
# ──────────────────────────────────────────────
def load_file(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file), uploaded_file.name
    elif name.endswith((".xlsx", ".xls")):
        return pd.read_excel(uploaded_file), uploaded_file.name
    elif name.endswith(".json"):
        data = json.load(uploaded_file)
        if isinstance(data, list):
            return pd.DataFrame(data), uploaded_file.name
        elif isinstance(data, dict):
            return pd.DataFrame([data]), uploaded_file.name
    raise ValueError(f"Unsupported file type: {uploaded_file.name}")


# ──────────────────────────────────────────────
#  EXPORT HELPERS
# ──────────────────────────────────────────────
def fig_to_bytes(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150,
                facecolor=fig.get_facecolor())
    buf.seek(0)
    return buf


def create_pdf_text_only(report: str) -> BytesIO:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 14)
    c.setFillColorRGB(0, 0.83, 1)
    c.drawString(40, height - 45, "BizIntel AI  —  Analysis Report")
    c.setFont("Helvetica", 9)
    c.setFillColorRGB(0.1, 0.1, 0.1)
    y = height - 70
    for line in report.split("\n"):
        wrapped = textwrap.wrap(line, 110) or [""]
        for wl in wrapped:
            c.drawString(40, y, wl)
            y -= 13
            if y < 50:
                c.showPage(); y = height - 50
                c.setFont("Helvetica", 9); c.setFillColorRGB(0.1, 0.1, 0.1)
    c.save(); buffer.seek(0); return buffer


def create_pdf_with_charts(report: str, figs: dict) -> BytesIO:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    c.setFont("Helvetica-Bold", 16)
    c.setFillColorRGB(0, 0.83, 1)
    c.drawString(40, height - 50, "BizIntel AI  —  Analysis Report")
    c.setFont("Helvetica", 9); c.setFillColorRGB(0.1, 0.1, 0.1)
    y = height - 75
    for line in report.split("\n"):
        wrapped = textwrap.wrap(line, 110) or [""]
        for wl in wrapped:
            c.drawString(40, y, wl); y -= 13
            if y < 50:
                c.showPage(); y = height - 50
                c.setFont("Helvetica", 9); c.setFillColorRGB(0.1, 0.1, 0.1)
    for name, fig in figs.items():
        if fig is None: continue
        c.showPage()
        c.setFont("Helvetica-Bold", 12); c.setFillColorRGB(0, 0.83, 1)
        c.drawString(40, height - 40, f"Chart: {name.replace('_', ' ').title()}")
        img_buf = fig_to_bytes(fig)
        img = ImageReader(img_buf)
        c.drawImage(img, 40, height - 460, width=width - 80, height=400, preserveAspectRatio=True)
    c.save(); buffer.seek(0); return buffer


def create_docx_text_only(report: str) -> BytesIO:
    doc = Document()
    heading = doc.add_heading("BizIntel AI — Analysis Report", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x00, 0x94, 0xFF)
    for line in report.split("\n"):
        p = doc.add_paragraph(line)
        if p.runs: p.runs[0].font.size = Pt(9)
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0); return buffer


def create_docx_with_charts(report: str, figs: dict) -> BytesIO:
    doc = Document()
    heading = doc.add_heading("BizIntel AI — Analysis Report", level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x00, 0x94, 0xFF)
    for line in report.split("\n"):
        p = doc.add_paragraph(line)
        if p.runs: p.runs[0].font.size = Pt(9)
    doc.add_page_break()
    doc.add_heading("Charts & Visualisations", level=2)
    for name, fig in figs.items():
        if fig is None: continue
        doc.add_heading(name.replace("_", " ").title(), level=3)
        img_buf = fig_to_bytes(fig)
        doc.add_picture(img_buf, width=Inches(5.5))
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0); return buffer


# ──────────────────────────────────────────────
#  LOGO
# ──────────────────────────────────────────────
def render_logo():
    try:
        c1, c2, c3 = st.columns([1, 1, 1])
        with c2:
            st.image("assets/BizIntel.png", width=220)
    except Exception:
        st.markdown("## 📊 BizIntel AI")


# ══════════════════════════════════════════════
#  AUTH PAGES (NOT LOGGED IN)
# ══════════════════════════════════════════════
if not st.session_state.logged_in:

    render_logo()

    # ─── LOGIN PAGE ───
    if st.session_state.page == "login":

        tab_login, tab_register = st.tabs(["🔑 Login", "📝 Register"])

        with tab_login:
            st.markdown("### Login to your account")
            email    = st.text_input("Email",    key="login_email")
            password = st.text_input("Password", type="password", key="login_pass")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Login", use_container_width=True):
                    success, user = login(email, password)
                    if success:
                        st.session_state.logged_in = True
                        st.session_state.user      = user
                        history = get_history(user[0])
                        st.session_state.first_visit = len(history) == 0
                        st.rerun()
                    else:
                        st.error("❌ Invalid email or password.")
            with col2:
                if st.button("Forgot Password", use_container_width=True):
                    st.session_state.page = "forgot"
                    st.rerun()

        with tab_register:
            st.markdown("### Create your account")
            r_email    = st.text_input("Email",     key="reg_email")
            r_name     = st.text_input("First Name", key="reg_name")
            r_surname  = st.text_input("Last Name",  key="reg_surname")
            r_phone    = st.text_input("Phone",      key="reg_phone")
            r_password = st.text_input("Password",   type="password", key="reg_pass")
            fav_language = st.selectbox(
                "Favorite Tech Language (used as security answer)",
                ["Python", "Java", "JavaScript", "C++", "C#", "PHP", "Go", "Rust", "Kotlin", "Swift"]
            )
            if st.button("Create Account", use_container_width=True):
                if not all([r_email, r_name, r_surname, r_phone, r_password, fav_language]):
                    st.warning("Please fill in all fields.")
                elif register(r_email, r_name, r_surname, r_phone, r_password, fav_language):
                    st.success("✅ Account created! Please log in.")
                else:
                    st.error("An account with that email already exists.")

    # ─── FORGOT PASSWORD PAGE ───
    elif st.session_state.page == "forgot":

        st.markdown("### 🔐 Reset Password")
        st.info("You have **2 attempts** to answer your security question. After 2 wrong answers you must wait **24 hours**.")

        identifier    = st.text_input("Email or Phone Number", key="fp_id")
        tech_answer   = st.text_input("Favorite Tech Language (your security answer)", key="fp_lang")
        new_password  = st.text_input("New Password", type="password", key="fp_newpass")
        new_password2 = st.text_input("Confirm New Password", type="password", key="fp_newpass2")

        col1, col2 = st.columns(2)

        with col1:
            if st.button("Reset Password", use_container_width=True):

                if not identifier or not tech_answer or not new_password:
                    st.warning("Please fill in all fields.")

                else:
                    # Check lockout
                    locked, secs_left = is_fp_locked(identifier)
                    if locked:
                        hrs  = secs_left // 3600
                        mins = (secs_left % 3600) // 60
                        st.error(
                            f"🔒 Too many failed attempts. "
                            f"Please try again in **{hrs}h {mins}m**."
                        )
                    else:
                        if new_password != new_password2:
                            st.error("New passwords do not match.")
                        elif len(new_password) < 6:
                            st.warning("Password must be at least 6 characters.")
                        else:
                            user = get_user(identifier)
                            if not user:
                                st.error("Account not found.")
                            elif user[5].lower().strip() != tech_answer.lower().strip():
                                attempts_used = record_fp_failure(identifier)
                                left = fp_attempts_left(identifier)
                                if left == 0:
                                    st.error(
                                        "❌ Wrong answer. You have used all 2 attempts. "
                                        "Please wait **24 hours** before trying again."
                                    )
                                else:
                                    st.error(
                                        f"❌ Wrong answer. You have **{left} attempt(s)** remaining."
                                    )
                            else:
                                reset_password(user[0], new_password)
                                reset_fp_attempts(identifier)
                                st.success("✅ Password reset successfully! You can now log in.")
                                time.sleep(1.5)
                                st.session_state.page = "login"
                                st.rerun()

        with col2:
            if st.button("Back to Login", use_container_width=True):
                st.session_state.page = "login"
                st.rerun()

    st.stop()


# ══════════════════════════════════════════════
#  MAIN APP (LOGGED IN)
# ══════════════════════════════════════════════
user    = st.session_state.user
email   = user[0]
name    = user[1]
surname = user[2]

# ── Sidebar ──
with st.sidebar:
    try:
        st.image("assets/BizIntel.png", width=160)
    except Exception:
        st.markdown("## 📊 BizIntel AI")

    st.markdown("---")
    st.markdown(f"**👤 {name} {surname}**")
    st.markdown(f"📧 {email}")
    st.markdown("---")

    # 🌙 Theme Toggle in sidebar
    theme_choice = st.radio(
        "🎨 Theme",
        ["Dark", "Light"],
        index=0 if st.session_state.theme == "Dark" else 1,
        horizontal=True
    )
    if theme_choice != st.session_state.theme:
        st.session_state.theme = theme_choice
        st.rerun()

    st.markdown("---")
    nav = st.radio("Navigation", ["📊 Dashboard", "📜 History", "⚙️ Settings"])
    st.markdown("---")
    if st.button("🚪 Logout", use_container_width=True):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()


greeting = "Welcome back" if not st.session_state.first_visit else "Welcome"
st.markdown(f"""
<div class="welcome-banner">
  <h2>👋 {greeting}, {name}!</h2>
  <p>BizIntel AI — Sentiment Analysis & Business Insights Platform</p>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════
#  TAB: DASHBOARD
# ══════════════════════════════════════════════
if "Dashboard" in nav:

    col_title, col_new = st.columns([5, 1])
    with col_new:
        if st.button("🔄 New Session", use_container_width=True,
                     help="Clear all data and start fresh"):
            st.session_state.df_analyzed     = None
            st.session_state.report_text     = ""
            st.session_state.summary         = None
            st.session_state.filename        = ""
            st.session_state.selected_column = None
            st.session_state.business_goal   = ""
            st.session_state.chat_history     = []
            st.session_state.all_figs         = {}
            st.rerun()

    st.markdown('<p class="section-header">📁 Upload Your Business Data</p>',
                unsafe_allow_html=True)
    st.caption("Supports CSV, Excel (.xlsx / .xls), and JSON")

    file = st.file_uploader(
        "Upload file",
        type=["csv", "xlsx", "xls", "json"],
        label_visibility="collapsed"
    )

    if file is not None:
        try:
            df_raw, fname = load_file(file)
        except Exception as e:
            st.error(f"Could not load file: {e}")
            st.stop()

        st.success(f"✅ Loaded **{fname}** — {df_raw.shape[0]:,} rows × {df_raw.shape[1]} columns")

        # Preview
        with st.expander("👁️ Preview Raw Data", expanded=False):
            st.dataframe(df_raw.head(20), use_container_width=True)
            st.caption(f"Columns: {', '.join(df_raw.columns.tolist())}")

        # ── Column selector (any column — text or numeric) ──
        all_columns = df_raw.columns.tolist()
        text_cols   = df_raw.select_dtypes(include=["object"]).columns.tolist()
        num_cols    = df_raw.select_dtypes(include=[np.number]).columns.tolist()

        col_a, col_b, col_c = st.columns(3)

        with col_a:
            selected_column = st.selectbox(
                "🎯 Column to Analyse",
                all_columns,
                help="Select any column — text, numeric, prices, ratings, etc."
            )

        with col_b:
            graph_types = st.multiselect(
                "📊 Chart Types",
                list(CHART_BUILDERS.keys()),
                default=["Bar Chart", "Pie Chart"]
            )

        with col_c:
            report_type = st.selectbox(
                "📋 Report Type",
                ["Quick Scan", "Executive Summary", "Full Report"],
                index=2
            )

        business_goal = st.text_area(
            "🎯 Business Goal / Context",
            value=st.session_state.business_goal,
            placeholder="Describe what you want to analyse (e.g. Identify top revenue drivers, find negative customer feedback patterns...)",
            height=80
        )
        st.session_state.business_goal = business_goal

        if st.button("🚀 Run Analysis", use_container_width=True):
            with st.spinner("Analysing data and computing sentiment..."):
                df = analyze_dataframe(df_raw.copy(), selected_column)
                summary = generate_summary(df)

                if report_type == "Quick Scan":
                    report = generate_quick_scan(df, selected_column, business_goal)
                elif report_type == "Executive Summary":
                    report = generate_executive_summary(df, selected_column, business_goal)
                else:
                    report = generate_full_report(df, selected_column, business_goal)

                st.session_state.df_analyzed     = df
                st.session_state.summary         = summary
                st.session_state.report_text     = report
                st.session_state.report_type     = report_type
                st.session_state.selected_column = selected_column
                st.session_state.filename        = fname
                st.session_state.graph_types     = graph_types
                st.session_state.first_visit     = False
                st.session_state.chat_history    = []
                st.session_state.all_figs        = {}

                save_history(
                    email, report, str(summary),
                    filename=fname,
                    total=summary["total"],
                    pos=summary["positive"],
                    neg=summary["negative"],
                    neu=summary["neutral"]
                )

            st.success("✅ Analysis complete!")
            st.rerun()

    # ── RESULTS SECTION ──
    if st.session_state.df_analyzed is not None:
        df      = st.session_state.df_analyzed
        summary = st.session_state.summary or {}
        report  = st.session_state.report_text
        col     = st.session_state.selected_column or df.columns[0]

        # ── Metric Cards ──
        st.markdown('<p class="section-header">📈 Key Metrics</p>', unsafe_allow_html=True)
        total = summary.get("total", 0)

        def metric_card(col_w, label, value, color):
            col_w.markdown(f"""
            <div class="metric-card">
                <div class="metric-value" style="color:{color}">{value}</div>
                <div class="metric-label">{label}</div>
            </div>""", unsafe_allow_html=True)

        m1, m2, m3, m4, m5, m6 = st.columns(6)
        metric_card(m1, "Total Records",     total,                         "#00d4ff")
        metric_card(m2, "✅ Positive",        summary.get("positive", 0),   "#22c55e")
        metric_card(m3, "❌ Negative",        summary.get("negative", 0),   "#ef4444")
        metric_card(m4, "➖ Neutral",          summary.get("neutral",  0),   "#f59e0b")
        metric_card(m5, "Avg Polarity",       summary.get("avg_polarity",  0), "#a78bfa")
        metric_card(m6, "Avg Subjectivity",   summary.get("avg_subjectivity", 0), "#38bdf8")

        # Numeric stats mini-table
        stats = numeric_stats(df)
        if stats:
            st.markdown('<p class="section-header">🔢 Numeric Column Statistics</p>', unsafe_allow_html=True)
            stats_df = pd.DataFrame(stats).T.reset_index().rename(columns={"index": "Column"})
            st.dataframe(stats_df, use_container_width=True, hide_index=True)

        # ── CHARTS ──
        st.markdown('<p class="section-header">📊 Visual Analysis</p>', unsafe_allow_html=True)

        for graph_type in st.session_state.graph_types:
            fig = None
            if graph_type == "Word Count Distribution":
                if col in df.select_dtypes(include=["object"]).columns:
                    fig = build_wordcount(df, col)
                else:
                    st.info(f"Word Count Distribution requires a text column (selected: '{col}' is numeric).")
                    continue
            elif graph_type == "Heatmap":
                fig = build_heatmap(df)
            elif graph_type == "Numeric Column Means":
                fig = build_numeric_chart(df)
            else:
                builder = CHART_BUILDERS.get(graph_type)
                if builder:
                    fig = builder(df)

            if fig:
                st.session_state.all_figs[graph_type] = fig
                st.subheader(graph_type)
                st.pyplot(fig)
                st.download_button(
                    f"📥 Download {graph_type}",
                    fig_to_bytes(fig),
                    f"{graph_type.lower().replace(' ', '_')}.png",
                    "image/png",
                    key=f"dl_chart_{graph_type}"
                )
            else:
                if graph_type in ("Heatmap", "Numeric Column Means"):
                    st.info(f"'{graph_type}' requires at least 2 numeric columns.")

        # ── REPORT ──
        st.markdown('<p class="section-header">📋 Insight Report</p>', unsafe_allow_html=True)
        st.markdown(f"**Active Report:** `{st.session_state.report_type}`")

        col_rt1, col_rt2, col_rt3 = st.columns(3)
        with col_rt1:
            if st.button("⚡ Quick Scan", use_container_width=True):
                with st.spinner("Generating Quick Scan..."):
                    r = generate_quick_scan(df, col, st.session_state.business_goal)
                st.session_state.report_text = r
                st.session_state.report_type = "Quick Scan"
                st.rerun()
        with col_rt2:
            if st.button("📝 Executive Summary", use_container_width=True):
                with st.spinner("Generating Executive Summary..."):
                    r = generate_executive_summary(df, col, st.session_state.business_goal)
                st.session_state.report_text = r
                st.session_state.report_type = "Executive Summary"
                st.rerun()
        with col_rt3:
            if st.button("📄 Full Report", use_container_width=True):
                extra = st.session_state.get("extra_sections_val", "")
                with st.spinner("Generating Full Report..."):
                    r = generate_full_report(df, col, st.session_state.business_goal, extra)
                st.session_state.report_text = r
                st.session_state.report_type = "Full Report"
                st.rerun()

        st.markdown(
            f'<div class="report-box">{report}</div>',
            unsafe_allow_html=True
        )

        # ── CHAT ──
        st.markdown('<p class="section-header">💬 Ask BizIntel AI Anything</p>', unsafe_allow_html=True)
        st.caption("e.g. 'What is the average revenue?', 'Show me the most negative records', 'What columns are missing values?'")

        for msg in st.session_state.chat_history:
            css_class = "chat-bubble-user" if msg["role"] == "user" else "chat-bubble-ai"
            icon      = "🧑" if msg["role"] == "user" else "🤖"
            st.markdown(
                f'<div class="{css_class}">{icon} {msg["content"]}</div>',
                unsafe_allow_html=True
            )

        user_q = st.text_input(
            "Your question",
            placeholder="e.g. What is the highest revenue? Which records are most negative?",
            label_visibility="collapsed",
            key="user_question_input"
        )
        qcol1, qcol2 = st.columns([3, 1])
        with qcol1:
            ask_btn = st.button("💬 Ask", use_container_width=True)
        with qcol2:
            if st.button("🗑️ Clear Chat", use_container_width=True):
                st.session_state.chat_history = []
                st.rerun()

        if ask_btn and user_q.strip():
            with st.spinner("Thinking..."):
                answer = answer_user_question(df, col, st.session_state.business_goal, user_q)
            st.session_state.chat_history.append({"role": "user",      "content": user_q})
            st.session_state.chat_history.append({"role": "assistant", "content": answer})
            st.rerun()

        # ── Analyzed Data Table ──
        with st.expander("📊 View Analyzed Dataset", expanded=False):
            st.dataframe(df, use_container_width=True)

        # ── DOWNLOADS ──
        st.markdown('<p class="section-header">⬇️ Download Report</p>', unsafe_allow_html=True)
        all_figs = st.session_state.all_figs

        dl1, dl2, dl3, dl4 = st.columns(4)
        dl1.download_button("📄 TXT Report",     create_pdf_text_only(report)  if False else report.encode(),
                            "bizintel_report.txt", "text/plain", use_container_width=True, key="dl_txt")
        dl2.download_button("📕 PDF (Text Only)", create_pdf_text_only(report),
                            "bizintel_report.pdf", "application/pdf", use_container_width=True, key="dl_pdf")
        dl3.download_button("📕 PDF + Charts",    create_pdf_with_charts(report, all_figs),
                            "bizintel_report_charts.pdf", "application/pdf", use_container_width=True, key="dl_pdf_c")
        dl4.download_button("📘 DOCX + Charts",   create_docx_with_charts(report, all_figs),
                            "bizintel_report_charts.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key="dl_docx_c")

        dl5, dl6 = st.columns(2)
        dl5.download_button("📘 DOCX (Text Only)", create_docx_text_only(report),
                            "bizintel_report.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key="dl_docx")
        csv_analyzed = df.to_csv(index=False).encode("utf-8")
        dl6.download_button("📊 CSV (Analyzed)", csv_analyzed, "analyzed_data.csv",
                            "text/csv", use_container_width=True, key="dl_csv")


# ══════════════════════════════════════════════
#  TAB: HISTORY
# ══════════════════════════════════════════════
elif "History" in nav:
    st.markdown('<p class="section-header">📜 Analysis History</p>', unsafe_allow_html=True)
    history = get_history(email)

    if not history:
        st.info("No analysis history yet. Run your first analysis on the Dashboard!")
    else:
        st.markdown(f"**{len(history)} analyses found**")
        for row in history:
            row_id, rpt_text, rpt_json, fname, total, pos, neg, neu, ts = (
                row[0], row[2], row[3], row[4], row[5], row[6], row[7], row[8], row[9]
            )
            pos_pct = round((pos / total) * 100, 1) if total else 0
            neg_pct = round((neg / total) * 100, 1) if total else 0
            neu_pct = round((neu / total) * 100, 1) if total else 0
            with st.expander(f"📁 {fname or 'Analysis'}  |  {ts[:16]}", expanded=False):
                st.markdown(f"""
                <div class="history-card">
                    <div class="history-meta">🕐 {ts} | File: {fname}</div><br>
                    <b>Records:</b> {total} &nbsp;|&nbsp;
                    <span style="color:#22c55e">✅ {pos} Positive ({pos_pct}%)</span> &nbsp;|&nbsp;
                    <span style="color:#ef4444">❌ {neg} Negative ({neg_pct}%)</span> &nbsp;|&nbsp;
                    <span style="color:#f59e0b">➖ {neu} Neutral ({neu_pct}%)</span>
                </div>
                """, unsafe_allow_html=True)
                st.text_area("Report", rpt_text, height=200, key=f"hist_{row_id}")
                hcol1, hcol2, hcol3 = st.columns(3)
                if hcol3.button(f"🗑️ Delete", key=f"del_{row_id}"):
                    from database import delete_history
                    delete_history(row_id)
                    st.success("Deleted.")
                    st.rerun()
                hcol1.download_button(f"📄 TXT", rpt_text, f"report_{row_id}.txt", key=f"dl_txt_{row_id}")
                hcol2.download_button(f"📕 PDF", create_pdf_text_only(rpt_text),
                                      f"report_{row_id}.pdf", "application/pdf", key=f"dl_pdf_{row_id}")


# ══════════════════════════════════════════════
#  TAB: SETTINGS
# ══════════════════════════════════════════════
elif "Settings" in nav:
    st.markdown('<p class="section-header">⚙️ Account Settings</p>', unsafe_allow_html=True)

    st.markdown(f"""
    <div class="history-card">
        <b>👤 Name:</b> {name} {surname}<br>
        <b>📧 Email:</b> {email}<br>
        <b>📱 Phone:</b> {user[3]}
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 🌙 Appearance")
    theme_sel = st.radio(
        "Theme",
        ["Dark", "Light"],
        index=0 if st.session_state.theme == "Dark" else 1,
        horizontal=True,
        key="settings_theme"
    )
    if theme_sel != st.session_state.theme:
        st.session_state.theme = theme_sel
        st.rerun()

    st.markdown("---")
    st.markdown("### 🔒 Change Password")
    curr_pass = st.text_input("Current Password", type="password")
    new_pass  = st.text_input("New Password",     type="password")
    conf_pass = st.text_input("Confirm New Password", type="password")

    if st.button("Update Password"):
        from auth import verify_password
        if not verify_password(curr_pass, user[4]):
            st.error("Current password is incorrect.")
        elif new_pass != conf_pass:
            st.error("New passwords do not match.")
        elif len(new_pass) < 6:
            st.warning("Password must be at least 6 characters.")
        else:
            reset_password(email, new_pass)
            st.success("✅ Password updated successfully.")

    st.markdown("---")
    st.markdown("### ⚠️ Danger Zone")
    confirm_delete = st.checkbox("I want to permanently delete my account and all data")
    if confirm_delete:
        if st.button("🗑️ Delete Account"):
            delete_user_account(email)
            st.session_state.clear()
            st.success("Account deleted.")
            st.rerun()

# This is needed for Streamlit Cloud deployment
if __name__ == "__main__":
    # The app will run automatically when streamlit run app.py is called
    pass